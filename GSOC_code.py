import os
import whisper
from docx import Document

audio_path = r"C:\Users\khana\Desktop\ASR_GSOC\audio\audio4.wav"

model_size = "small"
language = "Spanish"

output_docx = r"C:\Users\khana\Desktop\ASR_GSOC\audio4_transcription.docx"

if not os.path.exists(audio_path):
    print("Error: Audio file not found.")
    print("Please check the file path.")
else:
    try:
        print("Loading Whisper model...")
        model = whisper.load_model(model_size)

        print("Transcribing audio...")
        result = model.transcribe(audio_path, language=language)

        transcript = result["text"]

        print("\n===== TRANSCRIPTION OUTPUT =====\n")
        print(transcript)

        # Create Word document
        doc = Document()
        doc.add_heading("Audio Transcription", level=1)
        doc.add_paragraph(f"Audio File: {os.path.basename(audio_path)}")
        doc.add_paragraph(f"Language: {language}")
        doc.add_paragraph("")
        doc.add_paragraph(transcript)

        # Save document
        doc.save(output_docx)

        print(f"\nTranscription saved successfully to:\n{output_docx}")

    except Exception as e:
        print("An error occurred:")
        print(e)